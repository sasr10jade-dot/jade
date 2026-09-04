# -*- coding: utf-8 -*-
"""Generate VOICEMAP PRD v0.1 as a Word document from the finalized PRD content."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KOREAN_FONT = "맑은 고딕"
LATIN_FONT = "Calibri"

DARK = RGBColor(0x18, 0x18, 0x1B)
GRAY = RGBColor(0x71, 0x71, 0x7A)
LIGHT_GRAY = RGBColor(0xA1, 0xA1, 0xAA)
ACCENT = RGBColor(0x00, 0x00, 0x00)
AMBER = RGBColor(0x92, 0x40, 0x0E)
RED = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = LATIN_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_east_asian(run, font=KOREAN_FONT):
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def add_para(text="", size=10.5, bold=False, italic=False, color=None,
             space_after=6, space_before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
        set_east_asian(r)
    return p


def add_run_to(p, text, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    set_east_asian(r)
    return r


def add_section_heading(num, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}   {title}")
    r.font.size = Pt(17)
    r.bold = True
    r.font.color.rgb = DARK
    set_east_asian(r)
    add_para(desc, size=10, italic=False, color=GRAY, space_after=8)
    # divider
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    pPr = p2._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4D4D8")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_sub_heading(text, size=12, space_before=12):
    add_para(text, size=size, bold=True, color=DARK, space_before=space_before, space_after=4)


def add_bullets(items, size=10, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(it)
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        set_east_asian(r)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"


def add_table(headers, rows, widths=None, header_bg="27272A"):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_east_asian(r)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_bg)
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            set_east_asian(r)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(label, text, color=AMBER, bg="FEF3C7"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]) if isinstance(color, tuple) else "B45309")
    border.append(left)
    pPr.append(border)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = color
    set_east_asian(r1)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = color
    set_east_asian(r2)


# ============================================================
# COVER / HEADER
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("VOICEMAP")
r.font.size = Pt(28)
r.bold = True
set_east_asian(r)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Product Requirements Document")
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = GRAY
set_east_asian(r)

add_para("작곡가와 보컬을 연결하고, 구매와 저작권을 하나의 플로우로 관리하는 음악 마켓플레이스.",
         size=12, color=DARK, space_after=12)

meta_table = doc.add_table(rows=2, cols=4)
meta_table.style = "Light List Accent 1"
labels = ["PRD 버전", "상태", "업데이트", "작성자"]
values = ["v0.1", "Draft", "2026.05.13", "Product Team"]
for i, (lb, vl) in enumerate(zip(labels, values)):
    c0 = meta_table.rows[0].cells[i]
    c0.text = ""
    rr = c0.paragraphs[0].add_run(lb)
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = GRAY
    set_east_asian(rr)
    c1 = meta_table.rows[1].cells[i]
    c1.text = ""
    rr2 = c1.paragraphs[0].add_run(vl)
    rr2.font.size = Pt(10.5)
    rr2.bold = True
    set_east_asian(rr2)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

meta_table2 = doc.add_table(rows=2, cols=3)
meta_table2.style = "Light List Accent 1"
labels2 = ["PLATFORM", "TIMELINE", "OWNER"]
values2 = ["Web (Desktop First)", "MVP 8주 (2026 Q2)", "PM / Design / Eng"]
for i, (lb, vl) in enumerate(zip(labels2, values2)):
    c0 = meta_table2.rows[0].cells[i]
    c0.text = ""
    rr = c0.paragraphs[0].add_run(lb)
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = GRAY
    set_east_asian(rr)
    c1 = meta_table2.rows[1].cells[i]
    c1.text = ""
    rr2 = c1.paragraphs[0].add_run(vl)
    rr2.font.size = Pt(10.5)
    rr2.bold = True
    set_east_asian(rr2)

doc.add_page_break()

# ============================================================
# 01 Overview
# ============================================================
add_section_heading("01", "Overview", "왜 VOICEMAP을 만드는가")

add_sub_heading("Vision")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("소리를 잇는 다리.")
r.font.size = Pt(15)
r.bold = True
set_east_asian(r)
add_para("모든 곡에는 어울리는 목소리가 있다. VOICEMAP은 그 연결을 3분 안에 만든다.", color=GRAY, space_after=10)

add_sub_heading("Mission")
add_para("작곡가의 데모가 사장되지 않고, 보컬리스트가 오디션 없이 기회를 얻으며, 바이어가 저작권 걱정 없이 바로 구매하는 구조를 만든다.")

add_sub_heading("Problem")
add_para("현재 곡 거래는 카톡/DM, 가이드 녹음 비용 30~50만원, 정산·지분 구두 합의로 분쟁이 빈번함. 발견성 0, 신뢰성 0.")

add_sub_heading("Goals (MVP)")
add_bullets([
    "업로드 후 48시간 내 가이드 매칭",
    "구매 → 계약서 자동 생성 → 에스크로 결제",
    "곡당 Split 지분 시각화 및 합의 로그",
    "월 100건 거래 발생",
])

add_sub_heading("Non-Goals")
add_bullets([
    "DAW / 스트리밍 플레이어 내장",
    "AI 작곡 / AI 보컬 생성",
    "커뮤니티 / SNS 기능",
    "모바일 네이티브 앱 (MVP 이후)",
])

# ============================================================
# 02 Glossary
# ============================================================
add_section_heading("02", "Glossary", "용어 정의 - 팀 공통 언어")
add_table(
    ["용어", "영문", "정의"],
    [
        ["Creator", "작곡가/작사가", "트랙의 원작자. 음원 업로드 및 지분 설정 주체"],
        ["Performer", "보컬리스트", "가이드 보컬을 제공하고 실연권을 갖는 참여자"],
        ["Buyer", "바이어", "기획사 A&R, 아티스트, 콘텐츠 회사 등 구매자"],
        ["Track", "트랙", "원곡 음원 파일 (WAV/MP3) + 메타데이터"],
        ["Guide Vocal", "가이드 보컬", "곡의 느낌을 전달하기 위한 임시 보컬 녹음본"],
        ["Split", "지분", "저작권/수익 배분 비율. Creator-Performer 간 합의 (플랫폼 수수료는 별도 차감, Section 12 참고)"],
        ["Escrow", "에스크로", "구매 대금을 VOICEMAP이 보관 후 정산 조건 충족 시 분배"],
    ],
    widths=[3.2, 3.0, 8.5],
)

# ============================================================
# 03 Personas
# ============================================================
add_section_heading("03", "Personas", "누구를 위해 만드는가")

personas = [
    ("지훈", "작곡가 · 28세", "좋은 곡은 있는데 불러줄 사람이 없어요.",
     ["가이드 녹음 비용 부담", "DM으로 영업 피로", "지분 구두 합의 불안"], "월 2곡 이상 정식 발매 연결"),
    ("서아", "보컬리스트 · 24세", "제 목소리를 포트폴리오로 증명하고 싶어요.",
     ["오디션 정보 파편화", "무급 가이드 노동", "실연권 보장 없음"], "월 80만원 가이드 수익 + 크레딧 확보"),
    ("민수", "A&R · 기획사 · 35세", "저작권 정리된 곡을 빨리 찾고 싶어요.",
     ["곡 소싱 시간 과다", "지분 분쟁 리스크", "데모 퀄리티 편차"], "주 10곡 필터링 → 1곡 계약"),
]
for name, role, quote, pains, goal in personas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{name}  ")
    r.font.size = Pt(12); r.bold = True; set_east_asian(r)
    r2 = p.add_run(role)
    r2.font.size = Pt(9.5); r2.font.color.rgb = GRAY; set_east_asian(r2)
    add_para(f"“{quote}”", italic=True, color=DARK, space_after=4)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run("PAIN")
    r3.font.size = Pt(8.5); r3.bold = True; r3.font.color.rgb = GRAY; set_east_asian(r3)
    add_bullets(pains, size=9.5)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(2)
    r4 = p4.add_run("GOAL")
    r4.font.size = Pt(8.5); r4.bold = True; r4.font.color.rgb = GRAY; set_east_asian(r4)
    add_para(goal, size=9.5, bold=True, space_after=8)

# ============================================================
# 04 Core User Journey
# ============================================================
add_section_heading("04", "Core User Journey", "핵심 플로우 - 업로드부터 정산까지")

journey = [
    ("01", "Creator", "트랙 업로드", "P0", "WAV + 가사 + BPM/Key + 레퍼런스 태그"),
    ("02", "System", "보컬 매칭 큐", "P0", "장르/음역/분위기 기반 Performer 추천 · 알림"),
    ("03", "Performer", "가이드 보컬 제출", "P0", "가이드 녹음 업로드 + 실연권 동의 + 희망 Split 제안"),
    ("04", "Creator", "가이드 선택 & Split 확정", "P0", "최대 3개 가이드 비교 청취 → 최종 선택 + 지분 확정"),
    ("05", "Buyer", "디스커버리 & 청취", "P0", "필터(장르, 무드, BPM) → 가이드 포함 풀버전 청취"),
    ("06", "Buyer", "구매 & 에스크로 결제", "P0", "라이선스 선택(Exclusive/Non-Exclusive) → 결제"),
    ("07", "System", "계약 자동 생성", "P0", "Split 계약서 PDF + KOMCA 신고용 데이터 생성"),
    ("08", "System", "정산 분배", "P1", "플랫폼 수수료(15~20%, 라이선스별 차등) 제외 후 Creator/Performer 자동 분배"),
]
add_table(
    ["#", "주체", "단계", "우선순위", "설명"],
    journey,
    widths=[1.0, 2.2, 3.8, 1.8, 6.0],
)
add_callout("핵심 결정:", "가이드 보컬은 “구매 전환을 위한 필수 자산”이다. 가이드 없는 트랙은 디스커버리에서 노출 순위가 70% 낮아지도록 설계.")

# ============================================================
# 05 IA / Sitemap
# ============================================================
add_section_heading("05", "IA / Sitemap", "정보 구조")

add_sub_heading("Sitemap")
sitemap_lines = [
    "/",
    "├─ /discover — 탐색, 필터, 플레이어",
    "├─ /track/:id — 곡 상세, 가이드 비교, Split",
    "├─ /upload — Creator 업로드 플로우",
    "├─ /studio — 내 트랙, 가이드 관리",
    "├─ /inbox — 매칭 요청, 제안",
    "├─ /checkout/:id",
    "└─ /settings/split & /legal",
]
for line in sitemap_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    set_east_asian(r)

add_sub_heading("핵심 객체 모델")
obj_lines = [
    "Track { id, title, bpm, key, tags, stems? }",
    "Guide { performerId, audioUrl, splitAsk }",
    "Split { creator: 80%, performer: 20% }",
    "License { type, price, exclusive }",
    "Fee { rate: 15~20%, appliesTo: license.type }",
]
for line in obj_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    set_east_asian(r)

add_sub_heading("권한 매트릭스")
add_table(
    ["액션", "Creator", "Performer", "Buyer"],
    [
        ["업로드", "●", "-", "-"],
        ["가이드 제출", "-", "●", "-"],
        ["구매", "-", "-", "●"],
        ["Split 편집", "●", "○", "-"],
    ],
    widths=[4.0, 3.0, 3.0, 3.0],
)

# ============================================================
# 06 Functional Requirements
# ============================================================
add_section_heading("06", "Functional Requirements", "우선순위와 수용 기준")
fr_rows = [
    ["FR-01", "P0", "트랙 업로드 (WAV/MP3, 300MB)", "드래그&드롭, BPM/Key 자동 감지 실패 시 수동 입력, 업로드 진행률 표시"],
    ["FR-02", "P0", "가이드 보컬 매칭", "장르·음역·무드 태그 기반 추천 5명, 48시간 내 미응답 시 재매칭"],
    ["FR-03", "P0", "가이드 비교 플레이어", "동일 구간 A/B 전환, 파형 표시, 30초 프리뷰 제한 (바이어)"],
    ["FR-04", "P0", "Split 설정 및 합의", "Creator가 초기 지분 제안 → Performer 수락/역제안, 로그 타임라인 보관"],
    ["FR-05", "P0", "구매 & 라이선스 선택", "Exclusive / Non-Exclusive 가격 분리, 에스크로 결제 (토스페이먼츠)"],
    ["FR-06", "P0", "자동 계약서 생성", "PDF 생성, 전자서명(간편), KOMCA 신고용 CSV export"],
    ["FR-07", "P1", "정산 대시보드", "월별 수익, 곡별 판매, 수수료/Split 상세 내역, 예상 정산일 D-3 알림"],
    ["FR-08", "P1", "워터마크 & 다운로드 제어", "프리뷰 음원에 워터마크 삽입, 원본은 구매 후에만 다운로드"],
    ["FR-09", "P1", "리뷰 & 크레딧 시스템", "구매 후 Performer/Creator 상호 평가, 프로필에 누적 노출"],
    ["FR-10", "P2", "스탬프 트랙 & Stem 분리", "멀티트랙 업로드, Stem별 청취 토글 (향후 협업 확장)"],
]
add_table(["ID", "우선순위", "요구사항", "수용 기준 (AC)"], fr_rows, widths=[1.8, 1.8, 4.5, 6.7])

# ============================================================
# 07 Policy & Legal
# ============================================================
add_section_heading("07", "Policy & Legal", "신뢰를 위한 법적 장치")

policy = [
    ("저작권 (KOMCA)", [
        "업로드 시 원작자 서약 필수 (표절 아님)",
        "계약서에 KOMCA 등록 정보 포함",
        "Exclusive 판매 시 원저작권 이전 옵션 체크",
        "AI 생성물 업로드 금지 조항 (P0)",
    ]),
    ("실연권", [
        "가이드 보컬도 실연으로 간주, 실연권 20% 기본",
        "Performer 프로필에 실연자 정보 명시",
        "방송/공연 2차 사용 시 추가 정산 조항",
    ]),
    ("에스크로 & 정산", [
        "토스페이먼츠 에스크로: 구매 후 7일 이내 이의 없으면 정산",
        "플랫폼 수수료 Exclusive 20% / Non-Exclusive 15%, 시드 크리에이터 인하 요율 적용 (Section 12)",
        "월 1회 정산, 5만원 미만 이월",
        "환불: 원본 미다운로드 시 100%, 다운로드 후 50%",
    ]),
    ("분쟁 & 제재", [
        "표절 신고 시 72시간 내 블라인드 + 소명",
        "지분 분쟁 시 Split 로그가 Single Source of Truth",
        "3회 경고 시 업로드 제한",
        "법적 분쟁은 서울중앙지법 관할",
    ]),
]
for title, items in policy:
    add_sub_heading(title, size=11, space_before=8)
    add_bullets(items, size=9.5)

add_callout("법률 검토 필요 (외부 법무):",
            " 음악 저작권 양도 계약서 문구, 실연권 위임 범위, 에스크로 통신판매업 신고, 개인정보 처리방침(음원 파일의 개인정보 여부).")

# ============================================================
# 08 Metrics & Success Criteria
# ============================================================
add_section_heading("08", "Metrics & Success Criteria", "MVP 성공을 어떻게 측정하는가")

add_table(
    ["KPI", "목표", "설명"],
    [
        ["Time-to-Guide", "< 48h", "업로드 → 첫 가이드 제출까지"],
        ["Conversion", "> 12%", "청취 → 구매 전환율"],
        ["GMV", "₩20M / 월", "MVP 2개월 후 목표"],
    ],
    widths=[3.5, 2.5, 7.0],
)

add_table(
    ["지표", "정의", "MVP 목표", "측정 방법"],
    [
        ["매칭 성공률", "업로드 중 가이드 1개 이상 확보", "≥ 70%", "Track → Guide 이벤트"],
        ["Split 합의율", "제안 후 72h 내 합의", "≥ 85%", "Split status log"],
        ["NPS (Creator)", "재업로드 의향", "≥ 40", "구매 후 설문"],
        ["분쟁률", "거래 대비 분쟁 신고", "≤ 3%", "Support ticket"],
    ],
    widths=[2.8, 4.2, 2.0, 4.0],
)

# ============================================================
# 09 Roadmap
# ============================================================
add_section_heading("09", "Roadmap", "MVP 2개월 - 8주 플랜")

roadmap = [
    ("W1-2", "Foundation", "P0", ["Auth, 업로드, S3, 파형 플레이어", "Track 모델, 태그 시스템", "디자인 시스템 (Shadcn 기반)"]),
    ("W3-4", "Matching & Guide", "P0", ["Performer 온보딩, 매칭 알고리즘 v0 (룰베이스)", "가이드 업로드 & 비교 플레이어", "Inbox & 알림 (이메일)"]),
    ("W5-6", "Transaction", "P0", ["Split 에디터 & 합의 플로우", "토스페이먼츠 연동, 에스크로", "계약서 PDF 자동 생성 (Puppeteer)", "수수료 계산 엔진 (라이선스 유형별 차등 적용)"]),
    ("W7", "Discovery & Trust", "P1", ["Discover 페이지, 필터, 워터마크", "리뷰/크레딧 v0, 신고 기능"]),
    ("W8", "Launch Prep", "Launch", ["QA, 로드테스트, 법무 문구 확정", "랜딩 + 크리에이터 30명 시드", "ProductHunt & 트위터 론칭"]),
]
for week, title, status, items in roadmap:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{week}   {title}   [{status}]")
    r.font.size = Pt(11); r.bold = True; set_east_asian(r)
    add_bullets(items, size=9.5)

add_callout("Post-MVP (v0.2 - v1.0):",
            " Stem 협업, KOMCA API 연동, AI 가이드 매칭 고도화, 모바일 웹앱 최적화, 기획사 전용 워크스페이스. v1.0에서 누적 1,000 트랙, 월 GMV 1억 목표.",
            color=GRAY, bg="F4F4F5")

# ============================================================
# 10 Edge Cases & Error States
# ============================================================
add_section_heading("10", "Edge Cases & Error States", "주요 실패 시나리오와 처리 방식")
ec_rows = [
    ["EC-01", "P0", "업로드 실패 (포맷/용량 오류)", "지원 포맷 외 또는 300MB 초과 시 업로드 즉시 차단, 사유 명시 후 재시도 유도"],
    ["EC-02", "P1", "BPM/Key 자동 감지 실패", "신뢰도 낮음 시 수동 입력 폼으로 전환, 자동감지 배지 미표시"],
    ["EC-03", "P0", "매칭 무응답 (48h 초과)", "추천 5명 전원 무응답 시 자동 재매칭 큐 편입, Creator에게 알림"],
    ["EC-04", "P1", "가이드 중복 제출", "이미 확정된 트랙에 신규 제출 시 차단, 확정 완료 안내"],
    ["EC-05", "P0", "Split 합의 결렬", "역제안 3회 초과 시 합의 보류 상태로 전환, 양측에 조정 안내"],
    ["EC-06", "P0", "결제 실패", "PG 오류·한도초과 시 최대 3회 재시도, 실패 시 위시리스트 전환 안내"],
    ["EC-07", "P0", "에스크로 이의 제기", "구매 후 7일 내 이의 시 정산 자동 보류, 분쟁 프로세스로 라우팅"],
    ["EC-08", "P0", "Split 미확정 결제 시도", "결제 버튼 비활성화, Split 확정 필요 가드 노출"],
    ["EC-09", "P1", "워터마크 우회 시도", "프리뷰는 스트리밍 전용 재생, 다운로드 버튼 구매 전 비노출"],
    ["EC-10", "P1", "표절·저작권 신고 접수", "신고 3건 누적 시 즉시 블라인드 처리 + 72시간 내 소명 요청"],
]
add_table(["ID", "심각도", "케이스", "처리 방식"], ec_rows, widths=[1.8, 1.6, 4.2, 7.4])

# ============================================================
# 11 Open Questions & Risks
# ============================================================
add_section_heading("11", "Open Questions & Risks", "MVP 착수 전 확인이 필요한 이슈")

risks = [
    ("Product / Growth", [
        "콜드스타트: Creator·Performer 양측 동시 확보 필요 — 초기 시드 30명으로 공급 밀도 충분한가",
        "매칭 알고리즘 v0(룰베이스)의 초기 정확도 검증 방법 미정",
        "Non-Exclusive 재판매 시 가격 하한선 정책 미정",
        "BEP GMV(표준 요율 기준 약 ₩142M/월 추정)가 v1.0 목표(₩100M/월)를 상회 — 요율·고정비·목표 재검토 필요 (Section 12)",
    ]),
    ("Technical", [
        "300MB 오디오 업로드·트랜스코딩 처리 시간 및 인프라 비용 미검증",
        "매칭 큐 스케일링 전략 (동시 업로드 급증 시)",
        "워터마크 삽입 방식 (서버 사이드 vs 클라이언트) 미확정",
    ]),
    ("Legal / Compliance", [
        "KOMCA 공식 API 연동 일정 불확실 — MVP는 CSV export로 우회",
        "에스크로 운영을 위한 통신판매업 신고 소요 기간이 8주 로드맵과 상충 가능",
        "Exclusive 판매 시 원저작권 이전 범위의 법적 리스크 (법무 검토 대기)",
    ]),
    ("Open Decisions", [
        "수수료 체계(Section 12) 확정안 — 경쟁사 벤치마크 기반 최종 검증 필요",
        "모바일 대응 시점 (v0.2 vs v1.0)",
        "환불 기준(50%/100%)의 세부 예외 케이스 정의",
    ]),
]
for title, items in risks:
    add_sub_heading(title, size=11, space_before=8)
    add_bullets(items, size=9.5)

add_callout("우선 결정 필요 (Launch 전):",
            " 콜드스타트 공급 전략, 매칭 알고리즘 검증 방법, 통신판매업 신고 일정이 8주 로드맵의 핵심 리스크입니다.")

# ============================================================
# 12 Fee & Payout Structure
# ============================================================
add_section_heading("12", "Fee & Payout Structure", "수수료 체계와 정산 로직")

add_sub_heading("수수료 원칙")
add_bullets([
    "플랫폼 수수료는 구매 확정(에스크로 해제) 시점에 정가 기준으로 자동 차감",
    "PG(결제대행) 수수료는 플랫폼이 부담, 판매자·구매자에게 전가하지 않음",
    "Split 비율은 수수료 차감 후 판매자 몫(Net)에만 적용",
], size=9.5)

add_sub_heading("정산 주기")
add_bullets([
    "월 1회 정산, 5만원 미만은 다음 달로 이월",
    "정산 예정일 D-3 알림 (정산 대시보드 연동)",
    "분쟁·환불 진행 중인 건은 해당 정산 회차에서 자동 제외",
], size=9.5)

add_sub_heading("수수료율 (MVP)")
add_table(
    ["구분", "수수료율", "적용 대상", "비고"],
    [
        ["기본 (Exclusive)", "20%", "독점 라이선스 판매", "MVP 고정 요율"],
        ["기본 (Non-Exclusive)", "15%", "비독점 라이선스 판매", "반복 판매 가능, 건당 리스크 낮음"],
        ["시드 크리에이터 프로모션", "10%", "초기 시드 크리에이터 30명", "가입 후 90일 또는 5건 중 먼저 도달 시 종료"],
        ["볼륨 티어 (v0.2 예정)", "15~18%", "월 누적 판매액 300만원 이상", "MVP 이후 검증 후 도입, P2"],
    ],
    widths=[3.5, 2.0, 4.0, 5.0],
)

add_sub_heading("계산 예시 — Exclusive, 정가 ₩300,000")
calc_lines = [
    "구매 금액                 ₩300,000",
    "− 플랫폼 수수료 (20%)      ₩60,000",
    "= 판매자 정산액           ₩240,000",
    "  ├─ Creator (80%)        ₩192,000",
    "  └─ Performer (20%)       ₩48,000",
]
for line in calc_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    set_east_asian(r)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_sub_heading("환불 시 수수료 처리")
add_bullets([
    "원본 미다운로드 환불(100%): 플랫폼 수수료 전액 환불, 판매자 정산 취소",
    "다운로드 후 환불(50%): 플랫폼 수수료는 50%만 환불, 나머지는 플랫폼이 PG 실비로 흡수",
    "환불 발생 시 Split 정산은 자동 취소되며 Performer에게 별도 알림 발송",
], size=9.5)

add_sub_heading("플랫폼 마진 시뮬레이션 (추정)")
add_para("Metrics 목표 GMV ₩20,000,000/월 기준, PG 수수료 3.0%(카드 결제 업계 평균 추정) 차감 후 순마진",
         size=9, color=GRAY, space_after=6)
add_table(
    ["구간", "평균 수수료율", "총 수수료 수익", "PG 수수료(3%)", "순마진 / 순마진율"],
    [
        ["초기 (시드 프로모션 비중↑)", "12%", "₩2,400,000", "₩600,000", "₩1,800,000 (9.0%)"],
        ["표준 요율 정착 후", "18%", "₩3,600,000", "₩600,000", "₩3,000,000 (15.0%)"],
        ["볼륨 티어 도입 후 (v0.2)", "16%", "₩3,200,000", "₩600,000", "₩2,600,000 (13.0%)"],
    ],
    widths=[4.0, 2.2, 2.8, 2.6, 3.0],
)

add_sub_heading("손익분기점 (BEP) 추정")
add_para("4인 팀(PM 1 · Design 1 · Eng 2) 기준 고정비 가정, 순마진율로 필요 GMV 역산",
         size=9, color=GRAY, space_after=6)
add_table(
    ["고정비 항목", "월 비용"],
    [
        ["인건비 (4인, 평균 500만원)", "₩20,000,000"],
        ["인프라 (S3, CDN, 트랜스코딩)", "₩800,000"],
        ["기타 (툴, PG 기본료, CS)", "₩500,000"],
        ["합계", "₩21,300,000"],
    ],
    widths=[6.0, 4.0],
)
add_table(
    ["구간", "순마진율", "필요 BEP GMV/월"],
    [
        ["초기 (시드 프로모션)", "9.0%", "₩236,700,000"],
        ["표준 요율 정착 후", "15.0%", "₩142,000,000"],
        ["볼륨 티어 도입 후", "13.0%", "₩163,800,000"],
    ],
    widths=[4.5, 2.5, 4.0],
)

add_callout("갭 경고:",
            " 표준 요율 기준 BEP GMV(₩142,000,000/월)는 Post-MVP v1.0 목표(월 GMV ₩100,000,000)보다도 높습니다. "
            "MVP 목표(₩20M/월)는 BEP의 약 14% 수준 — 초기 적자는 예정된 경로이며, 흑자 전환에는 GMV 성장·고정비 최적화·요율 조정 중 "
            "최소 하나의 레버가 필요합니다.", color=RED, bg="FEE2E2")

add_callout("MVP 제안안:",
            " 본 수수료율은 초안이며, Section 11의 Open Decisions에 따라 경쟁사 벤치마크 검증 후 최종 확정됩니다. "
            "마진·BEP 시뮬레이션의 PG 수수료율(3.0%)과 고정비(인건비·인프라)는 업계 평균 기반 추정치이며, 실제 계약·채용 조건에 따라 "
            "변동될 수 있습니다 (Section 11 Technical Risk 참고).")

doc.add_page_break()

# ============================================================
# 13-15 Wireframes (text-form reference, low-fi)
# ============================================================
add_section_heading("13", "Dashboard Wireframe", "정산 대시보드 저해상도 와이어프레임 참고 (FR-07) — 실제 시각 디자인은 별도 제작")

add_sub_heading("화면 구성 요소")
add_bullets([
    "상단: 기간 선택 필터, CSV 내보내기 버튼",
    "KPI 3종: 이번 달 예상 정산액(₩2,940,000), 총 판매 건수(14건), 다음 정산일(D-3, 06.01)",
    "월별 수익 추이 (바 차트)",
    "정산 내역 테이블: 트랙명 / 판매일 / 판매액 / 수수료 / 내 몫 / 상태",
], size=9.5)
add_table(
    ["트랙명", "판매일", "판매액", "수수료", "내 몫", "상태"],
    [
        ["여름밤, 우리", "2026.05.02", "₩300,000", "₩60,000", "₩192,000", "정산완료"],
        ["Neon Drive", "2026.05.09", "₩150,000", "₩22,500", "₩102,000", "정산완료"],
        ["Waiting Room", "2026.05.14", "₩300,000", "₩60,000", "₩192,000", "진행중"],
    ],
    widths=[3.0, 2.2, 2.2, 2.0, 2.2, 2.0],
)

add_section_heading("14", "Core Screen Wireframes", "업로드 · Split 에디터 · 체크아웃 저해상도 와이어프레임 참고")

add_sub_heading("업로드 (FR-01)")
add_bullets([
    "드래그&드롭 업로드 영역 (WAV/MP3, 최대 300MB)",
    "제목 / BPM(자동감지) / Key(자동감지) 3종 필드",
    "장르·무드 태그 칩 (예: 발라드, 무드: 잔잔한, 여성보컬 추천)",
    "업로드 진행률 바",
    "액션: 임시저장 / 다음: 보컬 매칭",
], size=9.5)

add_sub_heading("Split 에디터 (FR-04)")
add_bullets([
    "트랙명 + 상태 배지(합의 대기중)",
    "Creator/Performer 지분 비율 바 (예: 80% / 20%)",
    "합의 로그: Creator 제안 → Performer 역제안 → Creator 재제안 타임라인",
    "액션: 역제안 / 수락",
], size=9.5)

add_sub_heading("체크아웃 (FR-05)")
add_bullets([
    "라이선스 선택 카드 2종: Exclusive(₩800,000, 독점·원저작권 이전 옵션) / Non-Exclusive(₩300,000, 비독점·반복 판매 가능)",
    "주문 요약: 트랙 / 라이선스 / 합계",
    "결제수단: 토스페이먼츠, 카드, 계좌이체",
    "에스크로 보호 안내 (구매 후 7일 이내 이의 없으면 정산, Section 7 참고)",
    "액션: 결제하기",
], size=9.5)

add_section_heading("15", "Additional Screen Wireframes", "Discover · 가이드 비교 플레이어 · Inbox · Studio 저해상도 와이어프레임 참고")

add_sub_heading("Discover (탐색)")
add_bullets([
    "필터 바: 장르 / 무드 / BPM / 정렬",
    "트랙 카드 리스트: 트랙명, 가격, BPM·Key·장르, 가이드 보컬 포함 여부",
    "예시: 여름밤, 우리(₩800,000) · Neon Drive(₩450,000) · Waiting Room(₩300,000)",
], size=9.5)

add_sub_heading("가이드 비교 플레이어 (FR-03)")
add_bullets([
    "가이드 A/B 파형 비교 플레이어 (동일 구간 A/B 전환)",
    "각 가이드의 Performer명 + 제안 Split (예: 서아 80/20, 민지 75/25)",
    "액션: A로 전환 / B로 전환 / 이 가이드로 확정",
], size=9.5)

add_sub_heading("Inbox — 매칭 요청함 (FR-02)")
add_bullets([
    "리스트형 알림함: 아바타, 이름, 메시지 미리보기, 시간, 안읽음 표시",
    "예시: 서아(가이드 제출), 민지(Split 역제안: 75/25), 민수 A&R(구매 문의)",
], size=9.5)

add_sub_heading("Studio — 내 트랙 관리 (/studio)")
add_bullets([
    "Creator의 트랙 리스트 + 상태 배지 (판매중 / Split 협의중 / 가이드 대기중)",
    "각 트랙의 가이드 확보 건수, 판매 건수 요약",
    "액션: + 새 트랙 업로드",
], size=9.5)
add_table(
    ["트랙명", "요약", "상태"],
    [
        ["여름밤, 우리", "가이드 2건 · 판매 3건", "판매중"],
        ["Neon Drive", "가이드 1건 · 판매 1건", "Split 협의중"],
        ["Waiting Room", "가이드 0건 · 판매 0건", "가이드 대기중"],
    ],
    widths=[3.5, 4.5, 3.0],
)

# ============================================================
# Footer
# ============================================================
doc.add_paragraph().paragraph_format.space_before = Pt(16)
add_para("© 2026 VOICEMAP · Draft · 내부 공유용 · 문의: product@voicemap.kr", size=9, color=GRAY)

out_path = "Voicemap-PRD-v0.1.docx"
doc.save(out_path)
print("saved:", out_path)
